"""Extract plain text from uploaded documents (TXT/MD/PDF/DOCX).

PDF/DOCX parsers are imported lazily so the rest of the KB works even if those
optional libraries aren't installed.
"""
from __future__ import annotations

import io


class UnsupportedDocument(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith((".txt", ".md", ".markdown", ".text")):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        return _pdf(data)
    if name.endswith(".docx"):
        return _docx(data)
    # Fall back to treating it as UTF-8 text.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        raise UnsupportedDocument(f"Unsupported document type: {filename!r}")


def _pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise UnsupportedDocument("PDF support requires the 'pypdf' package.")
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise UnsupportedDocument("DOCX support requires the 'python-docx' package.")
    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)
